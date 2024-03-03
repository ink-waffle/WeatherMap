import os
import docx
import folium
import math
import exifread
import datetime
import pandas as pd
import openmeteo_requests
import requests_cache
from retry_requests import retry
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement

# Функция для извлечения GPS-тегов из фото
def extract_gps_tags(image_path):
    with open(image_path, 'rb') as f:
        tags = exifread.process_file(f)

    lat_tag = tags.get('GPS GPSLatitude')
    lon_tag = tags.get('GPS GPSLongitude')
    date_taken_tag = tags.get('EXIF DateTimeOriginal')  # Добавляем тег для даты и времени создания фото
    
    if lat_tag and lon_tag:
        lat_values = [float(x.num) / float(x.den) if x.den != 0 else 0.0 for x in lat_tag.values]
        lon_values = [float(x.num) / float(x.den) if x.den != 0 else 0.0 for x in lon_tag.values]
        
        lat = lat_values[0] + lat_values[1] / 60 + lat_values[2] / 3600
        lon = lon_values[0] + lon_values[1] / 60 + lon_values[2] / 3600

        date_taken = datetime.datetime.strptime(str(date_taken_tag), '%Y:%m:%d %H:%M:%S') if date_taken_tag else None  # Преобразуем строку даты в объект datetime
        
        return lat, lon, date_taken  # Возвращаем широту, долготу и дату создания фото

    return None, None, None

# Функция для получения погодных данных
def get_weather_data(lat, lon, date_taken):
    session = requests_cache.CachedSession('.cache', expire_after = 3600)
    retry_session = retry(session, retries = 5, backoff_factor = 0.2)
    openmeteo = openmeteo_requests.Client(session = retry_session)
    
    url = "https://api.open-meteo.com/v1/forecast"
    params = {
        "latitude": lat,
        "longitude": lon,
        "hourly": "temperature_2m,precipitation,wind_speed_10m"
    }
    responses = openmeteo.weather_api(url, params=params)

    # Process first location. Add a for-loop for multiple locations or weather models
    response = responses[0]
    
    hourly = response.Hourly()
    hourly_temperature_2m = hourly.Variables(0).ValuesAsNumpy()
    hourly_precipitation = hourly.Variables(1).ValuesAsNumpy()
    hourly_wind_speed_10m = hourly.Variables(2).ValuesAsNumpy()

    hourly_data = {"date": pd.date_range(
        start=pd.to_datetime(hourly.Time(), unit="s", utc=True),
        end=pd.to_datetime(hourly.TimeEnd(), unit="s", utc=True),
        freq=pd.Timedelta(seconds=hourly.Interval()),
        inclusive="left"
    ), "temperature_2m": hourly_temperature_2m, "precipitation": hourly_precipitation,
        "wind_speed_10m": hourly_wind_speed_10m}

    hourly_dataframe = pd.DataFrame(data=hourly_data)
    return hourly_dataframe

# Функция для создания карты с маркерами
def create_coordinate_map(output_folder, image_folder):
    doc = docx.Document()
    doc.add_heading('Карта с GPS-тегами и погодой', 0)

    image_paths = [os.path.join(image_folder, filename) for filename in os.listdir(image_folder) if filename.lower().endswith(('.jpg', '.jpeg', '.tiff', '.png'))]
    print(image_paths)
    m = folium.Map(location=[0, 0], zoom_start=2)  # Начальное положение карты

    # Создаем таблицу в документе DOCX
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    
    # Настройки ширины столбцов
    table.autofit = False
    for width in [1, 3, 2, 2, 2, 2, 2, 2, 2, 2]:
        table.columns[width].width = Pt(100)
    
    # Заголовки столбцов
    table.rows[0].cells[0].text = '№'
    table.rows[0].cells[1].text = 'Широта-Долгота'
    table.rows[0].cells[2].text = 'Дата'
    table.rows[0].cells[3].text = 'Время'
    table.rows[0].cells[4].text = 'Осадки (мм)'
    table.rows[0].cells[5].text = 'Температура (°C)'
    table.rows[0].cells[6].text = 'Скорость ветра (км/ч)'
    table.rows[0].cells[7].text = 'Влажность (%)'
    
    for idx, image_path in enumerate(image_paths):
        lat, lon, date_taken = extract_gps_tags(image_path)

        if lat and lon and date_taken and all(math.isfinite(coord) for coord in (lat, lon)):
            lat, lon = round(lat, 6), round(lon, 6)
            date_str = date_taken.strftime('%Y-%m-%d')
            time_str = date_taken.strftime('%H:%M:%S')
            
            # Добавляем строку в таблицу
            row = table.add_row().cells
            row[0].text = str(idx + 1)
            row[1].text = f'{lat}, {lon}'
            row[2].text = date_str
            row[3].text = time_str

            weather_data = get_weather_data(lat, lon, date_taken)
            if 'precipitation' in weather_data.columns:
                precipitation = weather_data['precipitation']
                row[4].text = str(precipitation)
            else:
                precipitation = None
                row[4].text = 'Нет данных'
            if 'temperature_2m' in weather_data.columns:
                temperature_2m = weather_data['temperature_2m']
                row[5].text = str(temperature_2m)
            else:
                temperature_2m = None
                row[5].text = 'Нет данных'
            if 'wind_speed_10m' in weather_data.columns:
                wind_speed_10m = weather_data['wind_speed_10m']
                row[6].text = str(wind_speed_10m)
            else:
                wind_speed_10m = None
                row[6].text = 'Нет данных'
            if 'humidity_2m' in weather_data.columns:
                humidity_2m = weather_data['humidity_2m']
                row[7].text = str(humidity_2m)
            else:
                humidity_2m = None
                row[7].text = 'Нет данных'

            doc.add_paragraph(f'Фото: {image_path}', style='List Bullet')
            doc.add_paragraph(f'GPS-координаты: Широта {lat}, Долгота {lon}')
            doc.add_paragraph(f'Дата создания: {date_str}, Время создания: {time_str}')
            doc.add_paragraph(f'Осадки (мм): {precipitation}')
            doc.add_paragraph(f'Температура (°C): {temperature_2m}')
            doc.add_paragraph(f'Скорость ветра (км/ч): {wind_speed_10m}')
            doc.add_paragraph(f'Влажность (%): {humidity_2m}')
            doc.add_paragraph('', style='List Bullet')

            folium.Marker([lat, lon], popup=f'Фото: {image_path}\nОсадки: {precipitation} мм, Температура: {temperature_2m}°C, Скорость ветра: {wind_speed_10m} км/ч, Влажность: {humidity_2m} %').add_to(m)  # Добавление маркера на карту

    docx_file = os.path.join(output_folder, 'координаты_с_погодой.docx')
    doc.save(docx_file)
    messagebox.showinfo('Успех', f'Документ сохранен как {docx_file}')

    # Сохранение интерактивной карты в HTML файл
    map_output_file = os.path.join(output_folder, 'координаты_с_погодой.html')
    m.save(map_output_file)
    messagebox.showinfo('Успех', f'Интерактивная карта сохранена как {map_output_file}')

# Функция для обработки нажатия кнопки "Получить данные"
def get_data():
    image_folder = filedialog.askdirectory()
    if not image_folder:
        messagebox.showerror('Ошибка', 'Папка с фотографиями не выбрана.')
        return

    output_folder = filedialog.askdirectory()
    if not output_folder:
        messagebox.showerror('Ошибка', 'Папка для сохранения данных не выбрана.')
        return

    create_coordinate_map(output_folder, image_folder)

# Создание графического интерфейса
window = tk.Tk()
window.title('GPS Data Extractor')
window.geometry('400x200')  # Устанавливаем размер окна

# Создаем стиль для кнопки
style = ttk.Style()
style.configure('TButton', font=('Open Sans', 20), foreground='green', background='white', width=20, height=2)

get_data_button = ttk.Button(window, text='Получить данные', command=get_data, style='TButton')
get_data_button.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

window.mainloop()
